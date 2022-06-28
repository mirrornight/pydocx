import docx
from lxml import etree
import numpy as np

from pymdocx.common.comment import has_comment, parse_p_comment
from pymdocx.common.constants import WORD_NAMESPACE
from pymdocx.common.revision import has_revision


def get_label_in_paragraph(p, label_name):
    full_label_name = 'w:' + label_name
    xml = p._p.xml
    if full_label_name in xml:
        tree = etree.fromstring(xml)
        ins = tree.xpath('//' + full_label_name, namespaces=WORD_NAMESPACE)
    else:
        ins = []
    return ins


def print_xml_node(lxml_obj):
    print(etree.tostring(lxml_obj, pretty_print=True, encoding='utf-8').decode("utf-8"))


def print_doc(doc_file_path):
    document = docx.Document(docx=doc_file_path)
    print(document._body._element.xml)


def get_doc(doc_file_path):
    document = docx.Document(docx=doc_file_path)
    return document


def get_element_comment_revision_matrix(element):
    comment_list = []
    revision_list = []
    for p in element.paragraphs:
        comment_list.append(has_comment(p))
        revision_list.append(has_revision(p))
    return np.array([comment_list, revision_list])


def add_p_next(p1, p2):
    p1._p.addnext(p2._p)


def add_p_comment_next(p1, p2, comments_part_obj):
    # 添加段落p及其修订
    p1._p.addnext(p2._p)
    comment_list = parse_p_comment(p2)
    new_comment_id_list = []
    for comment in comment_list:
        comment_id = add_new_comment(comments_part_obj, comment['author'], comment['dtime'],
                                     comment['comment_text'],
                                     initials=comment['initials'])
        new_comment_id_list.append(comment_id)
    modify_oxml_element_comment_id(p2._p, new_comment_id_list)


def add_comment_2_p_end(p1, p2, comment_part_obj):
    comment_list = parse_p_comment(p2)
    for c in comment_list:
        p1.add_comment(c['comment_text'],
                       author=c['author'],
                       dtime=c['dtime'],
                       comment_part=comment_part_obj,
                       rangeStart=len(p1._p.getchildren()),
                       rangeEnd=len(p1._p.getchildren()))


def add_new_comment(comment_part, author, dtime, comment_text, initials=''):
    comment = comment_part.add_comment(author, initials, dtime)
    comment._add_p(comment_text)
    # CT_P _p
    # _r = p.add_r()
    # _r.add_comment_reference(comment._id)
    return comment._id


def modify_oxml_element_comment_id(oxml_element_obj, new_comment_id_list):
    attr_name = '{' + WORD_NAMESPACE['w'] + '}' + 'id'
    for i, cs in enumerate(oxml_element_obj.xpath('.//w:commentRangeStart')):
        cs.attrib[attr_name] = str(new_comment_id_list[i])
    for i, ce in enumerate(oxml_element_obj.xpath('.//w:commentRangeEnd')):
        ce.attrib[attr_name] = str(new_comment_id_list[i])
    for i, cr in enumerate(oxml_element_obj.xpath('.//w:commentReference')):
        cr.attrib[attr_name] = str(new_comment_id_list[i])


def _get_actual_p_index(has_add_mapping, doc_index, p_index):
    if doc_index in has_add_mapping.keys():
        has_add_mapping[doc_index] += 1
    else:
        has_add_mapping[doc_index] = 0
    return p_index - has_add_mapping[doc_index]